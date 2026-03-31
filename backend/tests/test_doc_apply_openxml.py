from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.doc_apply import apply_plan_to_docx, estimate_plan_impact
from app.schemas import Plan


def _xml_text(path: Path, entry: str) -> str:
    with zipfile.ZipFile(path, "r") as z:
        return z.read(entry).decode("utf-8")


def test_body_change_applies_east_asia_ascii_and_indent(tmp_path: Path) -> None:
    src = tmp_path / "src.docx"
    dst = tmp_path / "dst.docx"
    doc = Document()
    p = doc.add_paragraph("正文 Test 123")
    p.style = "Normal"
    doc.save(src)

    plan = Plan.model_validate(
        {
            "version": "1.0",
            "changes": [
                {
                    "scope": "body",
                    "font": {"east_asia": "宋体", "size_pt": 12},
                    "paragraph": {"first_line_indent_chars": 2, "line_spacing": "1.5"},
                }
            ],
        }
    )
    apply_plan_to_docx(src_docx=src, dst_docx=dst, plan=plan)
    xml = _xml_text(dst, "word/document.xml")
    assert "w:eastAsia=" in xml
    assert "Times New Roman" in xml
    assert "firstLineChars" in xml
    assert 'w:line="360"' in xml


def test_table_header_only_prefers_header_like_rows(tmp_path: Path) -> None:
    src = tmp_path / "src_tbl.docx"
    dst = tmp_path / "dst_tbl.docx"
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    for c in t.rows[0].cells:
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("表头")
        r.bold = True
    for c in t.rows[1].cells:
        c.paragraphs[0].add_run("数据")
    doc.save(src)

    plan = Plan.model_validate(
        {
            "version": "1.0",
            "changes": [{"scope": "tables_header_only", "font": {"east_asia": "黑体", "ascii": "Arial"}}],
        }
    )
    apply_plan_to_docx(src_docx=src, dst_docx=dst, plan=plan)
    xml = _xml_text(dst, "word/document.xml")
    # Header row text should have mutated fonts
    assert "黑体" in xml
    # Ensure second-row body text still exists (basic integrity check)
    assert "数据" in xml


def test_body_change_updates_header_footer_story_parts(tmp_path: Path) -> None:
    src = tmp_path / "src_hf.docx"
    dst = tmp_path / "dst_hf.docx"
    doc = Document()
    doc.add_paragraph("正文")
    sec = doc.sections[0]
    sec.header.paragraphs[0].text = "页眉 Header"
    sec.footer.paragraphs[0].text = "页脚 Footer"
    doc.save(src)

    plan = Plan.model_validate(
        {"version": "1.0", "changes": [{"scope": "body", "font": {"east_asia": "宋体", "ascii": "Times New Roman"}}]}
    )
    apply_plan_to_docx(src_docx=src, dst_docx=dst, plan=plan)

    with zipfile.ZipFile(dst, "r") as z:
        headers = [n for n in z.namelist() if n.startswith("word/header") and n.endswith(".xml")]
        footers = [n for n in z.namelist() if n.startswith("word/footer") and n.endswith(".xml")]
        assert headers and footers
        hx = z.read(headers[0]).decode("utf-8")
        fx = z.read(footers[0]).decode("utf-8")
    assert "w:rFonts" in hx
    assert "w:rFonts" in fx


def test_references_scope_uses_region_and_default_hanging_indent(tmp_path: Path) -> None:
    src = tmp_path / "src_ref.docx"
    dst = tmp_path / "dst_ref.docx"
    doc = Document()
    doc.add_paragraph("正文前言")
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] Author. Title.")
    doc.add_paragraph("[2] Author. Title 2.")
    doc.save(src)

    plan = Plan.model_validate(
        {
            "version": "1.0",
            "changes": [
                {
                    "scope": "references",
                    "font": {"east_asia": "宋体", "ascii": "Times New Roman", "size_pt": 12},
                    "paragraph": {"line_spacing": "single"},
                }
            ],
        }
    )
    apply_plan_to_docx(src_docx=src, dst_docx=dst, plan=plan)
    xml = _xml_text(dst, "word/document.xml")
    assert "hangingChars" in xml
    assert "Times New Roman" in xml


def test_references_scope_can_detect_trailing_reference_items_without_heading(tmp_path: Path) -> None:
    src = tmp_path / "src_ref2.docx"
    dst = tmp_path / "dst_ref2.docx"
    doc = Document()
    doc.add_paragraph("正文段落一")
    doc.add_paragraph("正文段落二")
    doc.add_paragraph("[1] Smith J. A paper title. 2020.")
    doc.add_paragraph("[2] Doe J. Another paper title. 2021.")
    doc.add_paragraph("[3] Foo B. Interesting article. 2022.")
    doc.save(src)

    plan = Plan.model_validate(
        {
            "version": "1.0",
            "changes": [{"scope": "references", "font": {"east_asia": "宋体", "ascii": "Times New Roman"}}],
        }
    )
    apply_plan_to_docx(src_docx=src, dst_docx=dst, plan=plan)
    xml = _xml_text(dst, "word/document.xml")
    assert "hangingChars" in xml


def test_references_scope_stops_before_appendix_block(tmp_path: Path) -> None:
    src = tmp_path / "src_ref3.docx"
    dst = tmp_path / "dst_ref3.docx"
    doc = Document()
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] AAA. 2020.")
    doc.add_paragraph("[2] BBB. 2021.")
    doc.add_paragraph("附录")
    doc.add_paragraph("这里是附录内容，不应当被按参考文献处理")
    doc.save(src)

    plan = Plan.model_validate(
        {
            "version": "1.0",
            "changes": [{"scope": "references", "font": {"east_asia": "宋体", "ascii": "Times New Roman"}}],
        }
    )
    apply_plan_to_docx(src_docx=src, dst_docx=dst, plan=plan)
    xml = _xml_text(dst, "word/document.xml")
    # Two reference paragraphs should get hanging indent, appendix should not.
    assert xml.count("hangingChars") == 2


def test_references_selector_can_limit_scan_window(tmp_path: Path) -> None:
    src = tmp_path / "src_ref4.docx"
    dst = tmp_path / "dst_ref4.docx"
    doc = Document()
    for _ in range(60):
        doc.add_paragraph("普通正文")
    doc.add_paragraph("[1] Long tail reference 2020.")
    doc.add_paragraph("[2] Long tail reference 2021.")
    doc.add_paragraph("[3] Long tail reference 2022.")
    doc.save(src)

    plan = Plan.model_validate(
        {
            "version": "1.0",
            "changes": [
                {
                    "scope": "references",
                    "selector": {"reference_max_scan_paragraphs": 20},
                    "font": {"east_asia": "宋体", "ascii": "Times New Roman"},
                }
            ],
        }
    )
    apply_plan_to_docx(src_docx=src, dst_docx=dst, plan=plan)
    xml = _xml_text(dst, "word/document.xml")
    assert "hangingChars" in xml


def test_estimate_plan_impact_counts_selected_paragraphs(tmp_path: Path) -> None:
    src = tmp_path / "src_impact.docx"
    doc = Document()
    doc.add_paragraph("A")
    doc.add_paragraph("B")
    doc.save(src)
    plan = Plan.model_validate({"version": "1.0", "changes": [{"scope": "body", "font": {"east_asia": "宋体"}}]})
    impact = estimate_plan_impact(src, plan)
    assert impact.get("body_paragraphs_changed_est") == 2

