from __future__ import annotations

from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path

from docx import Document

from .schemas import DocumentModelSummary, HeadingCandidate, StyleStat


UTC = timezone.utc


def analyze_docx(task_id: str, docx_path: Path) -> DocumentModelSummary:
    doc = Document(str(docx_path))

    paragraph_count = len(doc.paragraphs)
    table_count = len(doc.tables)

    # Style stats (lightweight, no full text)
    style_counter: Counter[str] = Counter()
    style_font_ea: dict[str, str] = {}
    style_font_ascii: dict[str, str] = {}
    style_font_size: dict[str, float] = {}
    style_bold_counts: dict[str, int] = defaultdict(int)
    style_center_counts: dict[str, int] = defaultdict(int)

    for p in doc.paragraphs:
        s = p.style.name if p.style is not None else "(None)"
        style_counter[s] += 1
        if p.alignment is not None and str(p.alignment).lower().endswith("center"):
            style_center_counts[s] += 1

        # Find first run with font info as sample
        for r in p.runs:
            if not style_font_ascii.get(s) and r.font and r.font.name:
                style_font_ascii[s] = r.font.name
            if not style_font_size.get(s) and r.font and r.font.size:
                try:
                    style_font_size[s] = float(r.font.size.pt)
                except Exception:
                    pass
            if r.bold:
                style_bold_counts[s] += 1
            # We cannot reliably extract eastAsia font from python-docx API; leave as None for now.
            if style_font_ascii.get(s) and style_font_size.get(s) is not None:
                break

    style_stats: list[StyleStat] = []
    for style_name, cnt in style_counter.most_common():
        bold_ratio = None
        if cnt > 0:
            # rough: bold run count / paragraph count (not perfect, but ok for summary)
            bold_ratio = min(1.0, style_bold_counts.get(style_name, 0) / max(1, cnt))
        center_ratio = None
        if cnt > 0:
            center_ratio = style_center_counts.get(style_name, 0) / cnt
        style_stats.append(
            StyleStat(
                style_name=style_name,
                paragraph_count=cnt,
                sample_font_east_asia=style_font_ea.get(style_name),
                sample_font_ascii=style_font_ascii.get(style_name),
                sample_font_size_pt=style_font_size.get(style_name),
                bold_ratio=bold_ratio,
                center_ratio=center_ratio,
            )
        )

    # Heading candidates (MVP: by style name only; later add heuristics scoring)
    heading_candidates: list[HeadingCandidate] = []
    for idx, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style is not None else ""
        level = None
        if style in {"标题 1", "Heading 1"}:
            level = 1
        elif style in {"标题 2", "Heading 2"}:
            level = 2
        elif style in {"标题 3", "Heading 3"}:
            level = 3
        if level:
            heading_candidates.append(
                HeadingCandidate(
                    paragraph_id=f"p{idx}",
                    style_name=style,
                    inferred_level=level,
                    confidence=0.9,
                )
            )

    # Sections / header/footer detection (lightweight)
    has_sections = len(doc.sections) > 1
    has_headers = any(bool(s.header and s.header.paragraphs) for s in doc.sections)
    has_footers = any(bool(s.footer and s.footer.paragraphs) for s in doc.sections)

    return DocumentModelSummary(
        task_id=task_id,
        created_at=datetime.now(tz=UTC),
        page_estimate=None,
        has_headers=has_headers,
        has_footers=has_footers,
        has_sections=has_sections,
        paragraph_count=paragraph_count,
        table_count=table_count,
        style_stats=style_stats,
        heading_candidates=heading_candidates,
    )

